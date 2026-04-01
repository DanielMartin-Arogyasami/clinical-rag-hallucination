"""Document parser — Stage 1. Handles PDF, DOCX, TXT, JSON."""
from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass
from dataclasses import field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    heading: str
    text: str
    level: int = 0
    page: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    document_id: str
    source_path: str
    format: str
    full_text: str
    sections: list[DocumentSection] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    tables: list[str] = field(default_factory=list)
    ocr_confidence: float | None = None


class DocumentParser:
    SUPPORTED_FORMATS = {"pdf", "docx", "txt", "json", "jsonl", "md"}

    def parse(self, path: str | Path) -> ParsedDocument:
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")
        ext = path.suffix.lstrip(".").lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: .{ext}")
        logger.info("Parsing %s (format=%s)", path.name, ext)
        return getattr(self, f"_parse_{ext}")(path)

    def _parse_txt(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(
            document_id=path.stem,
            source_path=str(path),
            format="txt",
            full_text=text,
            sections=self._segment_by_headings(text),
        )

    def _parse_md(self, path: Path) -> ParsedDocument:
        return self._parse_txt(path)

    def _parse_json(self, path: Path) -> ParsedDocument:
        raw = json.loads(path.read_text())
        if isinstance(raw, list):
            texts = [
                json.dumps(item, indent=2) if isinstance(item, dict) else str(item)
                for item in raw
            ]
            sections = [DocumentSection(heading=f"Record {i+1}", text=t) for i, t in enumerate(texts)]
            full_text = "\n\n".join(texts)
        elif isinstance(raw, dict):
            full_text = json.dumps(raw, indent=2)
            sections = [
                DocumentSection(
                    heading=k,
                    text=json.dumps(v, indent=2) if isinstance(v, (dict, list)) else str(v),
                )
                for k, v in raw.items()
            ]
        else:
            full_text = str(raw)
            sections = []
        return ParsedDocument(
            document_id=path.stem,
            source_path=str(path),
            format="json",
            full_text=full_text,
            sections=sections,
        )

    def _parse_jsonl(self, path: Path) -> ParsedDocument:
        records = [json.loads(line) for line in path.read_text().strip().split("\n") if line.strip()]
        sections = [
            DocumentSection(heading=f"Record {i+1}", text=json.dumps(r, indent=2))
            for i, r in enumerate(records)
        ]
        return ParsedDocument(
            document_id=path.stem,
            source_path=str(path),
            format="jsonl",
            full_text="\n\n".join(s.text for s in sections),
            sections=sections,
        )

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        full_text = "\n\n".join(pages)
        return ParsedDocument(
            document_id=path.stem,
            source_path=str(path),
            format="pdf",
            full_text=full_text,
            sections=self._segment_by_headings(full_text),
            metadata={"page_count": len(reader.pages)},
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        from docx import Document

        doc = Document(str(path))
        sections: list[DocumentSection] = []
        all_text: list[str] = []
        current_heading = "Document Start"
        current_text: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            all_text.append(text)
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                if current_text:
                    sections.append(
                        DocumentSection(heading=current_heading, text="\n".join(current_text))
                    )
                    current_text = []
                current_heading = text
            else:
                current_text.append(text)
        if current_text:
            sections.append(
                DocumentSection(heading=current_heading, text="\n".join(current_text))
            )
        tables = []
        for table in doc.tables:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            tables.append("\n".join(rows))
        return ParsedDocument(
            document_id=path.stem,
            source_path=str(path),
            format="docx",
            full_text="\n\n".join(all_text),
            sections=sections,
            tables=tables,
        )

    @staticmethod
    def _segment_by_headings(text: str) -> list[DocumentSection]:
        pattern = re.compile(r"^(?:\d+[\.\)]\s*)?[A-Z][A-Za-z\s:&\-]{2,80}$", re.MULTILINE)
        sections: list[DocumentSection] = []
        last_end = 0
        last_heading = "Document Start"
        for match in pattern.finditer(text):
            if match.start() > last_end:
                body = text[last_end : match.start()].strip()
                if body:
                    sections.append(DocumentSection(heading=last_heading, text=body))
            last_heading = match.group().strip()
            last_end = match.end()
        remaining = text[last_end:].strip()
        if remaining:
            sections.append(DocumentSection(heading=last_heading, text=remaining))
        return sections or [DocumentSection(heading="Full Document", text=text)]
